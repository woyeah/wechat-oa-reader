# SPDX-License-Identifier: AGPL-3.0-only
from __future__ import annotations

import asyncio
import logging
import re
from datetime import datetime
from io import BytesIO
from pathlib import Path

import httpx
from bs4 import BeautifulSoup
from bs4.element import NavigableString, Tag
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image

from .models import ArticleContent

LOGGER = logging.getLogger(__name__)

_BLOCK_TAGS = {
    "address",
    "article",
    "aside",
    "blockquote",
    "dd",
    "div",
    "dl",
    "dt",
    "figcaption",
    "figure",
    "footer",
    "form",
    "h1",
    "h2",
    "h3",
    "h4",
    "h5",
    "h6",
    "header",
    "hr",
    "li",
    "main",
    "nav",
    "ol",
    "p",
    "pre",
    "section",
    "table",
    "tbody",
    "td",
    "th",
    "thead",
    "tr",
    "ul",
}
_SKIP_TAGS = {"script", "style", "noscript"}
_SUPPORTED_DOCX_FORMATS = {"PNG", "JPEG", "JPG", "GIF", "BMP"}
_IMAGE_REQUEST_HEADERS = {
    "Referer": "https://mp.weixin.qq.com/",
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
        "(KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36"
    ),
}


def html_to_blocks(html: str) -> list[tuple[str, str]]:
    soup = BeautifulSoup(html, "html.parser")
    root: Tag | BeautifulSoup = soup.body or soup
    blocks: list[tuple[str, str]] = []
    text_buffer: list[str] = []

    def flush_text() -> None:
        if not text_buffer:
            return
        text = re.sub(r"\s+", " ", "".join(text_buffer)).strip()
        text_buffer.clear()
        if text:
            blocks.append(("text", text))

    def walk(node: Tag | NavigableString) -> None:
        if isinstance(node, NavigableString):
            text_buffer.append(str(node))
            return

        name = node.name.lower()
        if name in _SKIP_TAGS:
            return
        if name == "br":
            flush_text()
            return
        if name == "img":
            flush_text()
            image_url = (node.get("data-src") or node.get("src") or "").strip()
            if image_url:
                blocks.append(("image", image_url))
            return

        is_block = name in _BLOCK_TAGS
        if is_block:
            flush_text()
        for child in node.children:
            if isinstance(child, (Tag, NavigableString)):
                walk(child)
        if is_block:
            flush_text()

    for child in root.children:
        if isinstance(child, (Tag, NavigableString)):
            walk(child)
    flush_text()
    return blocks


async def article_to_docx(
    article: ArticleContent,
    output_path: Path,
    *,
    http_client: httpx.AsyncClient | None = None,
) -> Path:
    blocks = html_to_blocks(article.html)
    image_urls = [value for kind, value in blocks if kind == "image"]

    owns_client = http_client is None
    client = http_client or httpx.AsyncClient(follow_redirects=True)
    try:
        unique_urls = list(dict.fromkeys(image_urls))
        image_results = await asyncio.gather(*[_fetch_image(url, client) for url in unique_urls])
    finally:
        if owns_client:
            await client.aclose()

    image_map = dict(zip(unique_urls, image_results))

    document = Document()
    _configure_normal_style(document)
    document.add_heading(article.title or "", level=0)

    meta_parts: list[str] = []
    if article.author:
        meta_parts.append(article.author)
    publish_text = _format_publish_time(article.publish_time)
    if publish_text:
        meta_parts.append(publish_text)
    if meta_parts:
        meta_paragraph = document.add_paragraph(" · ".join(meta_parts))
        meta_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    url_paragraph = document.add_paragraph(article.url)
    url_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for kind, value in blocks:
        if kind == "text":
            document.add_paragraph(value)
            continue

        image_bytes = image_map.get(value)
        if image_bytes is None:
            LOGGER.warning("Failed to fetch image for docx export: %s", value)
            document.add_paragraph(f"[image failed: {value}]")
            continue

        try:
            width_in_inches = _calculate_display_width(image_bytes, max_width_inches=6.0)
            paragraph = document.add_paragraph()
            paragraph.add_run().add_picture(BytesIO(image_bytes), width=Inches(width_in_inches))
        except Exception:
            LOGGER.warning("Failed to embed image for docx export: %s", value, exc_info=True)
            document.add_paragraph(f"[image failed: {value}]")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path


async def _fetch_image(url: str, http_client: httpx.AsyncClient) -> bytes | None:
    try:
        response = await http_client.get(url, headers=_IMAGE_REQUEST_HEADERS)
        response.raise_for_status()
    except httpx.HTTPError:
        return None
    return _normalize_image_for_docx(response.content)


def _normalize_image_for_docx(image_bytes: bytes) -> bytes | None:
    try:
        with Image.open(BytesIO(image_bytes)) as image:
            image_format = (image.format or "").upper()
            if image_format in _SUPPORTED_DOCX_FORMATS:
                return image_bytes

            buffer = BytesIO()
            image.save(buffer, format="PNG")
            return buffer.getvalue()
    except Exception:
        return None


def _format_publish_time(publish_time: int) -> str | None:
    if publish_time <= 0:
        return None
    try:
        return datetime.fromtimestamp(publish_time).strftime("%Y-%m-%d %H:%M")
    except (OverflowError, OSError, ValueError):
        return None


def _configure_normal_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _calculate_display_width(image_bytes: bytes, *, max_width_inches: float) -> float:
    try:
        with Image.open(BytesIO(image_bytes)) as image:
            dpi_value = image.info.get("dpi", (96, 96))
            if isinstance(dpi_value, tuple) and dpi_value and dpi_value[0]:
                dpi = float(dpi_value[0])
            elif isinstance(dpi_value, (float, int)) and dpi_value:
                dpi = float(dpi_value)
            else:
                dpi = 96.0
            width_inches = image.width / dpi if dpi > 0 else max_width_inches
            if width_inches <= 0:
                return max_width_inches
            return min(width_inches, max_width_inches)
    except Exception:
        return max_width_inches
