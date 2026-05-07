# SPDX-License-Identifier: AGPL-3.0-only
from __future__ import annotations

from datetime import datetime
from io import BytesIO

import pytest
from docx import Document
from PIL import Image

from wechat_oa_reader.docx_writer import article_to_docx, html_to_blocks
from wechat_oa_reader.models import ArticleContent


def _tiny_png_bytes() -> bytes:
    buffer = BytesIO()
    Image.new("RGB", (8, 8), color=(255, 0, 0)).save(buffer, format="PNG")
    return buffer.getvalue()


def test_html_to_blocks_interleaves_text_and_images() -> None:
    html = """
    <div id="js_content">
      <p>First <strong>paragraph</strong></p>
      <section>
        Before image<br>After break
        <img data-src="https://img.example.com/a.webp" src="https://img.example.com/a.jpg" />
      </section>
      <p><img src="https://img.example.com/b.jpg" />Tail text</p>
    </div>
    """

    assert html_to_blocks(html) == [
        ("text", "First paragraph"),
        ("text", "Before image"),
        ("text", "After break"),
        ("image", "https://img.example.com/a.webp"),
        ("image", "https://img.example.com/b.jpg"),
        ("text", "Tail text"),
    ]


@pytest.mark.asyncio
async def test_article_to_docx_creates_doc_with_images(monkeypatch, tmp_path) -> None:
    image_bytes = _tiny_png_bytes()

    async def _fake_fetch_image(url: str, http_client) -> bytes | None:
        return image_bytes

    monkeypatch.setattr("wechat_oa_reader.docx_writer._fetch_image", _fake_fetch_image)

    article = ArticleContent(
        url="https://mp.weixin.qq.com/s/test",
        title="Docx Title",
        author="Author A",
        publish_time=1710000000,
        html="<div><p>Line one</p><p><img data-src='https://img.example.com/1.png' /></p><p>Line two</p></div>",
        plain_text="Line one\nLine two",
        images=["https://img.example.com/1.png"],
    )
    output_path = tmp_path / "article.docx"

    saved_path = await article_to_docx(article, output_path)

    assert saved_path == output_path
    assert output_path.exists()

    doc = Document(output_path)
    expected_time = datetime.fromtimestamp(article.publish_time).strftime("%Y-%m-%d %H:%M")

    assert doc.paragraphs[0].text == "Docx Title"
    assert doc.paragraphs[0].style.name == "Title"
    assert doc.paragraphs[1].text == f"Author A · {expected_time}"
    assert doc.paragraphs[2].text == article.url
    assert len(doc.inline_shapes) == 1

    paragraph_texts = [paragraph.text for paragraph in doc.paragraphs]
    assert "Line one" in paragraph_texts
    assert "Line two" in paragraph_texts


@pytest.mark.asyncio
async def test_article_to_docx_image_failure_inserts_placeholder(monkeypatch, tmp_path) -> None:
    image_bytes = _tiny_png_bytes()
    failed_url = "https://img.example.com/fail.png"

    async def _fake_fetch_image(url: str, http_client) -> bytes | None:
        if url == failed_url:
            return None
        return image_bytes

    monkeypatch.setattr("wechat_oa_reader.docx_writer._fetch_image", _fake_fetch_image)

    article = ArticleContent(
        url="https://mp.weixin.qq.com/s/test-failure",
        title="Failure Test",
        author="Author B",
        publish_time=0,
        html=(
            "<div><p>Intro</p>"
            "<p><img data-src='https://img.example.com/ok.png' /></p>"
            "<p><img data-src='https://img.example.com/fail.png' /></p></div>"
        ),
        plain_text="Intro",
        images=["https://img.example.com/ok.png", "https://img.example.com/fail.png"],
    )
    output_path = tmp_path / "article-failure.docx"

    saved_path = await article_to_docx(article, output_path)

    assert saved_path == output_path
    assert output_path.exists()

    doc = Document(output_path)
    assert len(doc.inline_shapes) == 1
    assert f"[image failed: {failed_url}]" in [paragraph.text for paragraph in doc.paragraphs]
